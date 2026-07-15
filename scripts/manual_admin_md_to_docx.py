#!/usr/bin/env python3
"""Convierte el manual Markdown a un .docx con diseño propio (sin dependencias externas, solo python-docx).

Paleta original (no el azul genérico de plantilla): tinta cálida + acento óxido/terracota,
apropiado para una empresa industrial de compresores. Tipografía serif en títulos, sans en
cuerpo. Control de "keep with next" en encabezados para que ningún título quede solo al pie
de una página. Imágenes reales de la app con marco sutil y pie de foto.

Uso: python scripts/manual_admin_md_to_docx.py docs/manual-panel-administrador.md docs/Manual-Panel-Administrador-CDV.docx
Las imágenes referenciadas en el markdown (![alt](manual-imagenes/x.png)) se resuelven
relativas a la carpeta del .md de origen.
"""
import os
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

SRC = sys.argv[1]
OUT = sys.argv[2]
SRC_DIR = os.path.dirname(os.path.abspath(SRC))

# ── Paleta ──────────────────────────────────────────────────────────────────
INK = RGBColor(0x24, 0x22, 0x1D)          # casi negro cálido, para títulos y cuerpo
ACCENT = RGBColor(0x9C, 0x53, 0x28)        # óxido / terracota, acento (no azul)
ACCENT_DARK = RGBColor(0x7A, 0x40, 0x1E)
MUTED = RGBColor(0x6E, 0x68, 0x5D)         # gris cálido para texto secundario
RULE = RGBColor(0xD9, 0xD2, 0xC2)          # regla fina, gris cálido claro
PAPER_EDGE = RGBColor(0xE9, 0xE4, 0xD8)

HEADING_FONT = "Cambria"
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
MD_LINK = re.compile(r"\[([^\]]+)\]\([^)]*\)")


def set_run_font(run, name):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_runs(paragraph, text, base_color=None):
    """Parsea **bold**, `code` y [links](#x) (sin el destino) en runs."""
    text = MD_LINK.sub(r"\1", text)
    pos = 0
    tokens = []
    pattern = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`")
    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append(("plain", text[pos:m.start()]))
        if m.group(1) is not None:
            tokens.append(("bold", m.group(1)))
        elif m.group(2) is not None:
            tokens.append(("code", m.group(2)))
        pos = m.end()
    if pos < len(text):
        tokens.append(("plain", text[pos:]))
    if not tokens:
        tokens = [("plain", text)]
    for kind, content in tokens:
        run = paragraph.add_run(content)
        if base_color is not None:
            run.font.color.rgb = base_color
        if kind == "bold":
            run.bold = True
            run.font.color.rgb = ACCENT_DARK if base_color is None else base_color
        elif kind == "code":
            set_run_font(run, MONO_FONT)
            run.font.size = Pt(9.5)
            run.font.color.rgb = ACCENT_DARK


def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    pPr.append(el)


def add_horizontal_rule(doc, color=RULE, weight=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(weight))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Compresores del Valle S.A.S. · Manual del Panel de Administrador · ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    set_run_font(run, BODY_FONT)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = p.add_run()
    page_run.font.size = Pt(8.5)
    page_run.font.color.rgb = MUTED
    page_run._element.append(fld_begin)
    page_run._element.append(instr)
    page_run._element.append(fld_end)


def add_image(doc, alt, rel_path):
    abs_path = os.path.normpath(os.path.join(SRC_DIR, rel_path))
    if not os.path.exists(abs_path):
        return
    with Image.open(abs_path) as im:
        w, h = im.size
    target_w_cm = 15.5
    target_h_cm = target_w_cm * (h / w)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(abs_path, width=Cm(target_w_cm), height=Cm(target_h_cm))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), "%02X%02X%02X" % (PAPER_EDGE[0], PAPER_EDGE[1], PAPER_EDGE[2]))
        pBdr.append(el)
    pPr.append(pBdr)
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        run_cap = cap.add_run(alt)
        run_cap.italic = True
        run_cap.font.size = Pt(9)
        run_cap.font.color.rgb = MUTED
        set_run_font(run_cap, BODY_FONT)


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.3)
        section.bottom_margin = Cm(2.1)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    heading_specs = {
        1: (21, ACCENT_DARK, 4, 4),
        2: (16, INK, 22, 3),
        3: (11.5, ACCENT_DARK, 12, 3),
    }
    for i, (size, color, before, after) in heading_specs.items():
        st = doc.styles[f"Heading {i}"]
        st.font.name = HEADING_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    list_bullet = doc.styles["List Bullet"]
    list_bullet.font.name = BODY_FONT
    list_bullet.font.size = Pt(10.5)
    list_bullet.font.color.rgb = INK
    list_bullet.paragraph_format.space_after = Pt(4)

    list_number = doc.styles["List Number"]
    list_number.font.name = BODY_FONT
    list_number.font.size = Pt(10.5)
    list_number.font.color.rgb = INK
    list_number.paragraph_format.space_after = Pt(4)

    lines = open(SRC, encoding="utf-8").read().splitlines()

    i = 0
    doc_title_done = False
    part_count = 0
    add_footer(doc.sections[0])

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        img_m = IMG_RE.match(stripped)
        if img_m:
            add_image(doc, img_m.group(1), img_m.group(2))
            i += 1
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if not doc_title_done:
                # Portada
                eyebrow = doc.add_paragraph()
                eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
                eyebrow.paragraph_format.space_before = Pt(150)
                r = eyebrow.add_run("COMPRESORES DEL VALLE S.A.S.")
                r.font.size = Pt(11)
                r.font.color.rgb = ACCENT
                r.bold = True
                set_run_font(r, BODY_FONT)

                title_p = doc.add_paragraph()
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_p.paragraph_format.space_before = Pt(14)
                title_p.paragraph_format.space_after = Pt(0)
                parts = text.split("·")
                r = title_p.add_run(parts[0].strip())
                r.bold = True
                r.font.size = Pt(30)
                r.font.color.rgb = INK
                set_run_font(r, HEADING_FONT)

                rule_p = doc.add_paragraph()
                rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rule_p.paragraph_format.space_before = Pt(16)
                pPr = rule_p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "%02X%02X%02X" % (ACCENT[0], ACCENT[1], ACCENT[2]))
                pBdr.append(bottom)
                pPr.append(pBdr)
                rr = rule_p.add_run(" ")
                rr.font.size = Pt(1)

                sub = doc.add_paragraph()
                sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub.paragraph_format.space_before = Pt(18)
                r = sub.add_run("Guía completa del panel administrativo, pantalla por pantalla")
                r.italic = True
                r.font.size = Pt(12.5)
                r.font.color.rgb = MUTED
                set_run_font(r, HEADING_FONT)

                doc.add_page_break()
                doc_title_done = True
            else:
                part_count += 1
                doc.add_page_break()
                eyebrow = doc.add_paragraph()
                eyebrow.paragraph_format.space_before = Pt(60)
                eyebrow.paragraph_format.space_after = Pt(2)
                r = eyebrow.add_run(f"PARTE {part_count}")
                r.font.size = Pt(11)
                r.bold = True
                r.font.color.rgb = ACCENT
                set_run_font(r, BODY_FONT)
                title_text = text.split(".", 1)[-1].strip() if "." in text[:3] else text
                h = doc.add_paragraph()
                keep_with_next(h)
                h.paragraph_format.space_after = Pt(6)
                r = h.add_run(title_text)
                r.bold = True
                r.font.size = Pt(24)
                r.font.color.rgb = INK
                set_run_font(r, HEADING_FONT)
                keep_with_next(add_horizontal_rule(doc, color=ACCENT, weight=10))
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            if text == "Índice":
                text = "Contenido"
            h = doc.add_paragraph(style="Heading 2")
            add_runs(h, text, base_color=INK)
            keep_with_next(add_horizontal_rule(doc))
            i += 1
            continue

        if stripped.startswith("### "):
            text = stripped[4:].strip()
            h = doc.add_paragraph(style="Heading 3")
            add_runs(h, text.upper(), base_color=ACCENT_DARK)
            i += 1
            continue

        if stripped.startswith("_") and stripped.endswith("_") and len(stripped) > 2:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(stripped[1:-1])
            run.italic = True
            run.font.color.rgb = MUTED
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:].strip())
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            content = stripped.lstrip(">").strip()
            add_runs(p, content)
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = MUTED
            i += 1
            continue

        # párrafo normal
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
