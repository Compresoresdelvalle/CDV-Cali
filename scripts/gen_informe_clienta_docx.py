# -*- coding: utf-8 -*-
"""
Genera el informe para la clienta en Word (.docx) con un diseño sobrio y
editorial: portada limpia con logo, secciones con antetitulo + linea fina,
numeros de punto discretos, mucho espacio en blanco y casi sin color
(tinta, gris y un acento muy puntual de azul marino).

Lenguaje sencillo, sin guiones largos, facil de entender.

Uso:  python scripts/gen_informe_clienta_docx.py
Salida: docs/Informe-Mejoras-CDV-2026-06-28.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta sobria: tinta, gris y un acento minimo de azul marino ─────────────
INK = RGBColor(0x1A, 0x1D, 0x21)
NAVY = RGBColor(0x14, 0x36, 0x5C)
GRAY = RGBColor(0x70, 0x76, 0x80)
MUTE = RGBColor(0xA6, 0xAC, 0xB4)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

NAVY_HEX = "14365C"
LINE_HEX = "D8DBDF"     # lineas finas
HEAD_HEX = "EFF1F3"     # encabezado de tabla (gris muy claro)
ALT_HEX = "FAFBFC"      # franja de tabla apenas perceptible
NOTE_HEX = "F6F7F9"     # nota (casi blanco)

LOGO = os.path.join("src", "assets", "logo-cdv.jpeg")

doc = Document()

for s in doc.sections:
    s.top_margin = Inches(0.95)
    s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)
    s.different_first_page_header_footer = True

base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(11)
base.paragraph_format.space_after = Pt(6)
base.paragraph_format.line_spacing = 1.12

# Orden de hijos de w:pPr (para insertar shd/pBdr en su lugar del esquema).
PPR_ORDER = [
    "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore", "w:framePr",
    "w:widowControl", "w:numPr", "w:suppressLineNumbers", "w:pBdr", "w:shd",
    "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
]


def add_pPr_child(pPr, tagname):
    el = OxmlElement(tagname)
    successors = PPR_ORDER[PPR_ORDER.index(tagname) + 1:]
    pPr.insert_element_before(el, *successors)
    return el


def set_run(run, font="Calibri", size=11, color=INK, bold=False, italic=False,
            spacing=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if spacing is not None:
        rPr = run._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rPr.append(sp)


def shade_para(p, fill):
    shd = add_pPr_child(p._p.get_or_add_pPr(), "w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def para_borders(p, **edges):
    pbdr = add_pPr_child(p._p.get_or_add_pPr(), "w:pBdr")
    for edge, (color, sz, space) in edges.items():
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), str(space))
        e.set(qn("w:color"), color)
        pbdr.append(e)


def cell_fill(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def table_borders(t, color=LINE_HEX, sz=4):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    for edge in ("left", "right", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    t._tbl.tblPr.insert_element_before(
        borders, "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook",
        "w:tblCaption", "w:tblDescription", "w:tblPrChange",
    )


def add_page_field(paragraph):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(instr); run._r.append(e)
    return run


# ── Bloques de diseño ────────────────────────────────────────────────────────
def spacer(pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p


def part(kicker, title):
    pk = doc.add_paragraph()
    pk.paragraph_format.space_before = Pt(22)
    pk.paragraph_format.space_after = Pt(1)
    r = pk.add_run(kicker.upper())
    set_run(r, "Calibri", 9, NAVY, bold=True, spacing=44)

    pt = doc.add_paragraph()
    pt.paragraph_format.space_after = Pt(9)
    para_borders(pt, bottom=(LINE_HEX, 6, 7))
    r = pt.add_run(title)
    set_run(r, "Georgia", 16, INK, bold=True)
    return pt


def item(num, title):
    label = f"{num:02d}" if isinstance(num, int) else str(num)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + "    ")
    set_run(r1, "Georgia", 12.5, MUTE, bold=True)
    r2 = p.add_run(title)
    set_run(r2, "Georgia", 12.5, INK, bold=True)
    return p


def lead(texto, etiqueta=None):
    par = doc.add_paragraph()
    if etiqueta:
        r = par.add_run(etiqueta.rstrip(":").upper() + "   ")
        set_run(r, "Calibri", 9, INK, bold=True, spacing=20)
    par.add_run(texto)
    return par


def bullet(texto):
    return doc.add_paragraph(texto, style="List Bullet")


def nota(texto, etiqueta="Para tener en cuenta:"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.right_indent = Pt(8)
    shade_para(p, NOTE_HEX)
    para_borders(p, left=(NAVY_HEX, 14, 12))
    r = p.add_run(" " + etiqueta + " ")
    set_run(r, "Calibri", 10, NAVY, bold=True, italic=True)
    r2 = p.add_run(texto)
    set_run(r2, "Calibri", 10, GRAY)
    return p


def tabla(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_borders(t)
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        cell_fill(hdr[i], HEAD_HEX)
        run = hdr[i].paragraphs[0].add_run(htext)
        set_run(run, "Calibri", 9.5, INK, bold=True)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                cell_fill(cells[i], ALT_HEX)
            run = cells[i].paragraphs[0].add_run(str(val))
            set_run(run, "Calibri", 9.5, INK if i == 0 else GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def build_footer():
    fp = doc.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_borders(fp, top=(LINE_HEX, 6, 6))
    r = fp.add_run("Compresores del Valle S.A.S.       Informe de mejoras y correcciones       ")
    set_run(r, "Calibri", 8.5, GRAY)
    pr = add_page_field(fp)
    set_run(pr, "Calibri", 8.5, GRAY)


# ── Portada (limpia y centrada) ──────────────────────────────────────────────
def portada():
    spacer(46)
    pl = doc.add_paragraph()
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.add_run().add_picture(LOGO, width=Inches(1.35))
    pl.paragraph_format.space_after = Pt(10)

    pn = doc.add_paragraph()
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pn.add_run("COMPRESORES DEL VALLE S.A.S.")
    set_run(r, "Calibri", 11, GRAY, bold=True, spacing=50)

    # Linea fina corta y centrada (unico acento de color)
    pr = doc.add_paragraph()
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr.paragraph_format.left_indent = Inches(2.7)
    pr.paragraph_format.right_indent = Inches(2.7)
    pr.paragraph_format.space_before = Pt(14)
    pr.paragraph_format.space_after = Pt(14)
    para_borders(pr, bottom=(NAVY_HEX, 8, 2))

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pt.add_run("Informe de mejoras y correcciones")
    set_run(r, "Georgia", 26, INK, bold=True)
    pt.paragraph_format.space_after = Pt(2)

    psub = doc.add_paragraph()
    psub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = psub.add_run("Todo lo que se mejoro y se corrigio en la aplicacion, explicado facil")
    set_run(r, "Georgia", 12.5, GRAY, italic=True)

    pdte = doc.add_paragraph()
    pdte.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pdte.paragraph_format.space_before = Pt(4)
    r = pdte.add_run("28 de junio de 2026")
    set_run(r, "Calibri", 10.5, GRAY, spacing=12)

    spacer(34)
    nota(
        "por cada punto encontraras tres partes muy cortas: que pasaba Antes, "
        "que se hace Ahora, y En que ayuda. Al final hay un pequeno glosario con "
        "las palabras que pueden sonar raras.",
        etiqueta="Como leer este informe:",
    )
    doc.add_page_break()


# ═════════════════════════════ CONTENIDO ═════════════════════════════════════
portada()

lead(
    "Revisamos a fondo como la aplicacion cuenta el dinero (el cierre de caja, "
    "el conteo del efectivo, las cuentas por cobrar y por pagar, y como se "
    "relacionan con las reparaciones). Encontramos y corregimos varios casos en "
    "los que el sistema contaba el mismo dinero dos veces, o lo contaba antes de "
    "tiempo. Tambien agregamos varias mejoras que pediste y dejamos la "
    "aplicacion mucho mas comoda de usar desde el celular."
)
nota(
    "el conteo del efectivo de la caja ya venia funcionando bien. Los errores "
    "afectaban sobre todo los totales del cierre (Ingresos, Egresos y Margen) y "
    "el tablero de resumen, que ahora quedan iguales al efectivo real.",
    etiqueta="Importante:",
)

# Parte 1
part("Parte 1", "El cierre de caja y el dinero")

item(1, "El dinero de las reparaciones se contaba dos veces")
lead("cuando un cliente pagaba una reparacion en abonos y luego se le entregaba el equipo, el sistema sumaba ese dinero dos veces: una vez cuando el cliente abonaba y otra vez cuando se entregaba y facturaba. Si el abono y la entrega caian en cierres distintos, el mismo dinero aparecia en dos cierres y los ingresos se veian mas altos de lo real.", "Antes:")
lead("una reparacion recibio un abono de $400.000 el 16 de junio y se entrego el 24 de junio. El sistema volvia a sumar la venta el 24, asi que el mismo dinero quedaba contado en dos cierres.", "Ejemplo real:")
lead("el dinero de una reparacion se cuenta una sola vez: el dia en que el cliente abona. La factura ya no se vuelve a sumar (es solo el comprobante; el dinero ya habia entrado con los abonos).", "Ahora:")
lead("los ingresos por servicios reflejan el dinero que de verdad entro, sin repetir.", "En que ayuda:")

item(2, "El efectivo de una reparacion entregada el mismo dia desaparecia del conteo")
lead("si una reparacion se pagaba en efectivo y se entregaba el mismo dia, ese efectivo no aparecia en el efectivo esperado de la caja. La caja mostraba un sobrante sin explicacion al cuadrar.", "Antes:")
lead("el efectivo de los abonos se cuenta siempre por su fecha, aunque la reparacion ya se haya entregado. La caja vuelve a cuadrar con el efectivo real.", "Ahora:")
lead("se elimina un descuadre que aparecia al entregar reparaciones el mismo dia del pago.", "En que ayuda:")

item(3, "Los cobros y pagos de cartera no llegaban al conteo de caja")
lead("cuando se cobraba una venta a credito en efectivo, o se pagaba a un proveedor en efectivo, ese movimiento no se reflejaba en el efectivo esperado de la caja.", "Antes:")
lead("la caja suma los cobros en efectivo y resta los pagos en efectivo. El efectivo esperado ya incluye todos los movimientos reales del dia.", "Ahora:")
lead("el conteo de caja tiene en cuenta tambien lo que entra y sale por creditos.", "En que ayuda:")

item(4, "Las compras y ventas a credito afectaban el cierre sin haberse pagado")
lead("al registrar una compra a credito, el sistema la dejaba bien en Cuentas por Pagar, pero al mismo tiempo la contaba como gasto en el cierre ese mismo dia, aunque todavia no se hubiera pagado nada. Lo mismo pasaba con las ventas a credito: se contaban como ingreso el dia de la venta, aunque el cliente todavia no hubiera pagado.", "Antes:")
lead("ahora el credito solo afecta el cierre cuando el dinero de verdad se mueve:", "Ahora:")
tabla(
    ["Operacion", "Antes", "Ahora"],
    [
        ["Compra a credito", "Gasto el dia de la factura", "Solo en Cuentas por Pagar; entra al cierre el dia que se paga"],
        ["Venta a credito", "Ingreso el dia de la venta", "Solo en Cuentas por Cobrar; entra al cierre el dia que se cobra"],
        ["Compra de caja menor", "Gasto (correcto)", "Gasto (igual, sin cambios)"],
    ],
)
lead("una compra a credito de $251.506 del 27 de junio ya no aparece como gasto ese dia (queda en $0 hasta que se pague). Cuando se registre el pago, ahi si entra al cierre, el dia del pago.", "Ejemplo real:")
lead("Ingresos, Egresos y Margen reflejan el dinero real. El credito vive en cartera hasta que se paga o se cobra.", "En que ayuda:")

item(5, "El tablero de resumen mostraba ingresos de servicios mas altos de lo real")
lead("el tablero del administrador calculaba los ingresos de servicios del mes con el mismo doble conteo del punto 1.", "Antes:")
lead("se corrigio igual que el cierre: el dinero de las reparaciones se cuenta una sola vez, por su fecha de abono.", "Ahora:")
lead("el tablero muestra cifras de servicios que coinciden con el cierre.", "En que ayuda:")

# Parte 2
part("Parte 2", "Cuentas por cobrar, por pagar y reparaciones")

item(6, "Anular un cobro o pago ahora pide confirmacion y motivo")
lead("en la pantalla de cobros y pagos, al tocar el boton de eliminar, el movimiento se anulaba al instante, sin preguntar ni guardar el motivo. Habia riesgo de anular por error.", "Antes:")
lead("al anular un cobro o pago, el sistema ahora:", "Ahora:")
bullet("Pide confirmacion.")
bullet("Permite escribir el motivo de la anulacion.")
bullet("Guarda el movimiento en el historial (no se borra; queda como anulado con quien, cuando y por que).")
lead("se evita borrar por error y queda registro de cada anulacion.", "En que ayuda:")

item(7, "La forma de pago correcta al pasar una cotizacion a venta")
lead("al convertir una cotizacion aprobada en venta, quedaba siempre como efectivo, aunque tuviera saldo pendiente. Asi una venta con saldo aparecia como si ya estuviera pagada.", "Antes:")
lead("ahora la venta toma la forma de pago real:", "Ahora:")
bullet("Si la cotizacion ya estaba pagada, la venta queda como Efectivo.")
bullet("Si queda saldo pendiente, la venta queda como Credito y entra a Cuentas por Cobrar.")
lead("las ventas que vienen de cotizacion muestran su estado de pago verdadero.", "En que ayuda:")

item(8, "Control de abonos en reparaciones sin valor")
lead("se podia registrar un abono de cualquier monto en una reparacion que todavia no tenia valor (total en $0). Eso permitia anticipos enormes (por ejemplo, un abono de $400.000 en una orden que despues valia $2.000) que luego descuadraban el cierre.", "Antes:")
lead("el sistema pide cotizar primero la reparacion (ponerle el valor de repuestos y mano de obra) antes de aceptar abonos. El asistente guia en ese orden: primero la cotizacion, luego el anticipo.", "Ahora:")
lead("se evita registrar anticipos sobre ordenes sin valor, una fuente de descuadres.", "En que ayuda:")

# Parte 3
part("Parte 3", "Mejoras de uso")

item(9, "Nueva seccion para los equipos que se arman")
lead("no habia forma de administrar desde la aplicacion la lista de equipos que se pueden armar.", "Antes:")
lead("en Configuracion se agrego la seccion de equipos armables. Desde ahi el administrador puede:", "Ahora:")
bullet("Crear un equipo (incluso con un nombre provisional, para confirmarlo despues).")
bullet("Editar su nombre, referencia y precio.")
bullet("Quitar un equipo de la lista.")
nota("quitar solo lo saca de la lista de armables; no afecta su venta ni su inventario en el resto de la aplicacion.")
lead("el administrador maneja esa lista sin depender de soporte tecnico.", "En que ayuda:")

item(10, "Filtro por fecha en los listados")
lead("en los listados de Ventas, Cotizaciones, Reparaciones y Traspasos no se podia filtrar por fecha; tocaba bajar a mano.", "Antes:")
lead("en la misma barra de busqueda que ya tenian, ahora se puede escribir una fecha:", "Ahora:")
bullet("Si escribes dia, mes y ano (por ejemplo 15/06/2026), muestra los documentos de ese dia.")
bullet("Si escribes mes y ano (por ejemplo 06/2026), muestra los de ese mes completo.")
lead("Si no escribes una fecha, la busqueda funciona como siempre (por numero, cliente y demas).")
lead("encontrar documentos de un dia o un mes en segundos, sin pantallas nuevas.", "En que ayuda:")

# Parte 4
part("Parte 4", "Permisos y busqueda")

item(11, "La devolucion de herramientas queda en Bodega y Administracion")
lead("cualquier usuario con acceso a Herramientas (incluidos los tecnicos) podia registrar la devolucion de una herramienta. El prestamo si estaba limitado, pero la devolucion no.", "Antes:")
lead("la devolucion de herramientas queda solo para Bodega y Administracion. Los tecnicos ya no ven esa accion.", "Ahora:")
tabla(
    ["Accion sobre herramientas", "Quien puede"],
    [
        ["Crear y prestar", "Bodega, Administracion"],
        ["Devolver", "Bodega, Administracion (antes: cualquiera de la sede)"],
        ["Regresar a inventario de insumo", "Solo Administracion"],
    ],
)
lead("el control de entrada y salida de herramientas queda en manos de Bodega y Administracion, como debe ser.", "En que ayuda:")

item(12, "Buscar referencias con punto decimal (por ejemplo 2.5 o 3.5)")
lead("la barra de busqueda quitaba el punto. Por eso, al buscar una polea 2.5 o 3.5, el sistema buscaba 25 o 35 y no encontraba el producto. Pasaba con cualquier referencia que tuviera un punto.", "Antes:")
lead("la busqueda respeta el punto. Escribir 2.5 encuentra los productos cuya referencia contiene 2.5.", "Ahora:")
lead("se pueden encontrar productos por referencias con decimales, como las poleas.", "En que ayuda:")

# Parte 5
part("Parte 5", "Reparaciones y cambios de producto")

item(13, "Una reparacion no autorizada ya no obliga a sacar repuestos")
lead("cuando se cotizaban repuestos y luego el cliente NO autorizaba la reparacion, el sistema no dejaba cerrar la orden a menos que se sacaran del inventario esos repuestos. Es decir, obligaba a sacar piezas que no se iban a usar, solo para poder cobrar la revision.", "Antes:")
lead("si el cliente no autoriza, la orden se cierra cobrando solo la revision o diagnostico, sin tocar el inventario:", "Ahora:")
bullet("Ya no aparece el boton de sacar del inventario en ese caso.")
bullet("Una nota explica que los repuestos cotizados no se sacan porque no hubo autorizacion.")
bullet("Al marcar como terminada, la cotizacion de repuestos se descarta sola, sin afectar el inventario.")
lead("una orden no autorizada se cierra rapido y bien, cobrando solo la revision, sin gastar repuestos que el cliente no aprobo.", "En que ayuda:")

item(14, "Nuevo: cambio de producto con cobro o devolucion de la diferencia")
lead("que cuando un cliente regrese a cambiar un producto por otro de distinto precio, el sistema maneje la diferencia: cobrar si el nuevo es mas caro o devolver si es mas barato, todo ligado a la factura original.", "Lo que se pidio:")
lead("se agrego el boton Registrar cambio dentro de la factura original (no hay que buscar la venta ni teclear numeros). El paso a paso es:", "Ahora:")
bullet("Se elige, de la factura, que producto devuelve el cliente y la cantidad.")
bullet("Se busca y elige el producto nuevo que se lleva y su cantidad.")
bullet("El sistema calcula solo la diferencia y muestra si hay que cobrar o devolver, y cuanto.")
lead("Al confirmar: el producto devuelto vuelve al inventario, el nuevo sale del inventario, y si el nuevo es mas caro se cobra solo la diferencia (efectivo o transferencia); si es mas barato, se devuelve la diferencia al cliente en efectivo.")
lead("el cierre cuenta solo la diferencia (lo que de verdad entro o salio), una sola vez. El cambio queda ligado a la venta original para poder rastrearlo.", "En el dinero:")
lead("los cambios se hacen en un solo paso, con el inventario y la caja siempre cuadrados.", "En que ayuda:")
nota("en esta primera version, la diferencia se cobra en efectivo o transferencia, y cuando hay que devolver al cliente se hace en efectivo. La opcion de nota de credito (saldo a favor) para una compra futura se puede agregar mas adelante como mejora aparte.")
nota("para deshacer un cambio no se anula la venta de la diferencia (eso descuadraria el inventario). Se hace el cambio al reves: se devuelve el producto nuevo y se entrega de vuelta el original. El sistema muestra una nota recordandolo en esa venta.", etiqueta="Como deshacer un cambio:")

# Parte 6
part("Parte 6", "Roles y permisos")

item(15, "Los tecnicos ya no operan ni facturan reparaciones (solo se asignan)")
lead("un tecnico podia operar toda la reparacion (diagnostico, cotizacion, autorizacion, sacar repuestos e incluso facturar y entregar), igual que Ventas. Eso mezclaba responsabilidades.", "Antes:")
lead("ahora estan separados:", "Ahora:")
bullet("Tecnicos: quedan disponibles para asignarse a una reparacion y la ven solo para consultar. Ya no pueden crear, avanzar ni facturar; tampoco les aparece el boton de nueva orden.")
bullet("Ventas y Administracion: operan todo el proceso (recepcion, diagnostico, autorizacion, repuestos, facturacion y cierre).")
lead("esto se controla tanto en la pantalla como por dentro del sistema, para que no se pueda saltar.")
lead("vista mas simple para los tecnicos y la facturacion en el perfil correcto.", "En que ayuda:")

item(16, "El perfil de Bodega puede ver el cierre para consultar")
lead("un perfil para quien recibe caja, que pueda ver el cierre y el inventario de todas las bodegas, sin poder configurar, borrar ni cambiar cosas importantes.", "Lo que se pidio:")
lead("ahora:", "Ahora:")
bullet("El perfil Bodega tiene en su menu la opcion Cierre (solo para ver): elige fechas, ve el resumen de totales y consulta el historial. No puede generar ni firmar cierres (eso sigue siendo solo de Administracion).")
bullet("Ver inventario de todas las bodegas: ya estaba disponible para Bodega.")
bullet("No tiene acceso al Panel de Administracion: se mantiene restringido.")
lead("quien maneja la caja puede cuadrar y revisar el cierre del dia sin poder alterar nada, desde el celular o el computador.", "En que ayuda:")

# Parte 7
part("Parte 7", "Pago mixto")
item(17, "Nuevo: pago combinado (efectivo y transferencia) en una sola factura")
lead("que cuando un cliente pague combinando efectivo y transferencia, el sistema permita registrar las dos formas en la misma venta, validando que la suma sea igual al total, y que el cierre las muestre por separado.", "Lo que se pidio:")
lead("en Nueva venta, ademas de Efectivo, Transferencia, Tarjeta y Credito, hay una opcion Mixto:", "Ahora:")
bullet("Se escribe el monto en efectivo y el monto por transferencia (con su cuenta).")
bullet("El sistema muestra en vivo la suma y el total, y solo deja confirmar cuando coinciden exactamente.")
bullet("Al guardar, cada forma de pago queda registrada en su medio.")
lead("la venta cuenta una sola vez en los ingresos (sin doble conteo), pero el dinero se reparte: el efectivo entra a la caja (y al conteo de efectivo, que espera solo esa parte) y la transferencia entra a su cuenta bancaria.", "En el cierre:")
lead("se pueden cobrar facturas con pago combinado sin inventar dos ventas, y la caja del dia cuadra con el efectivo real.", "En que ayuda:")

# Parte 8
part("Parte 8", "Revision final de calidad")
lead("despues de implementar todo lo anterior se hizo una segunda revision del manejo del dinero. No se encontraron errores que afectaran el dia a dia, pero se afinaron tres detalles pequenos para dejarlos perfectos.")

item(18, "Cambio de producto sobre una factura con descuento")
lead("si la factura original tenia un descuento general, el credito por el producto devuelto se calculaba sobre el precio de lista (sin el descuento), acreditando de mas. Ahora el credito refleja lo que el cliente de verdad pago. Las facturas sin descuento no cambian en nada.")

item(19, "El detalle por producto del cierre, ya descontando el descuento")
lead("en el detalle por producto del cierre, el ingreso de cada producto ahora descuenta su parte del descuento. Antes mostraba el precio de lista, lo que inflaba el reporte de cualquier venta con descuento. Es solo el detalle informativo: los totales de caja (Ingresos, Egresos, Margen y conteo) no cambian.")

item(20, "El limite de cantidad en la pantalla de cambio")
lead("la pantalla de cambio ahora resta lo que ya se habia devuelto de esa misma factura, para no permitir pedir mas unidades de las que quedan (muestra el mensaje quedan N). El sistema ya lo impedia por seguridad; ahora tambien se ve reflejado en la pantalla.")

# Parte 9
part("Parte 9", "Uso desde el celular")
lead("se hizo una revision completa del uso desde el celular (como se navega, que tan comodos son los botones con guantes y que ve cada tipo de usuario). El hallazgo principal: en el celular la barra de abajo solo tenia 5 botones y era la unica forma de moverse, asi que cada usuario solo alcanzaba entre el 30% y el 50% de sus secciones. Se rediseno la navegacion y se corrigieron varios detalles.")

item(21, "Acceso completo a todas las secciones desde el celular")
lead("en el celular solo se veian 5 botones abajo; el resto de las secciones del usuario quedaban inalcanzables.", "Antes:")
lead("se agrego un menu Mas (se abre con el boton de abajo a la derecha o tocando el avatar de arriba) que muestra todas las secciones del usuario, ordenadas por grupo, mas el buscador. La barra de abajo conserva los accesos mas usados con un boton central destacado (Vender, Orden o Inventario, segun el usuario). La cobertura paso del 30%-50% al 100%.", "Ahora:")
lead("desde el celular se llega a todo lo que el usuario puede hacer, no solo a 5 atajos.", "En que ayuda:")

item(22, "El Panel de administracion completo en el celular")
lead("en el celular el Panel de administracion solo mostraba 5 de sus 12 herramientas; las otras 7 no se podian abrir.", "Antes:")
lead("el Panel tiene su propia barra abajo (Resumen, Alertas, Conteo, Auditoria y Mas) y un menu Mas con las 12 herramientas y una salida clara a Operaciones. Todo se alcanza desde el celular.", "Ahora:")
lead("la administracion se puede operar completa desde el telefono.", "En que ayuda:")

item(23, "Botones mas grandes para usar con guantes")
lead("en la pantalla de venta (la mas usada) los botones de cantidad y precio eran muy pequenos y dificiles de tocar con guantes; el cambio de producto igual. Ademas, en una pantalla, la barra de accion de abajo podia quedar tapada por el menu inferior en celulares con el borde redondeado.", "Antes:")
lead("los botones de cantidad y precio y los del cambio se hicieron mas grandes (tamano comodo para guantes), y la barra de accion ahora respeta el borde del celular para no taparse con el menu de abajo.", "Ahora:")
lead("menos errores al tocar y nada queda oculto en celulares modernos.", "En que ayuda:")

item(24, "Bodega ya puede ver el resumen del cierre antes de generarlo")
lead("el perfil Bodega ya tenia la opcion Cierre (solo para ver, punto 16), pero al pedir el resumen el sistema respondia con un error (decia que solo el administrador podia consultar cierres). Quedaba a medias.", "Antes:")
lead("Bodega puede ver el resumen del cierre y el historial sin problemas. Generar o firmar el cierre sigue siendo solo de Administracion.", "Ahora:")
lead("quien maneja la caja puede cuadrar el dia de verdad, sin poder cambiar nada.", "En que ayuda:")

item(25, "Ajustes del menu y los permisos por tipo de usuario")
bullet("Tecnico: se quito Productos de su menu (aparecia, pero la pantalla lo rechazaba y lo devolvia al inicio). Su menu queda en Ordenes, Ensambles y Herramientas, acorde con su rol.")
bullet("Abonos de reparaciones: en linea con que el tecnico solo consulta, el registro de abonos queda en Ventas y Administracion (la pantalla ya lo hacia; se reforzo tambien por dentro).")
lead("cada usuario ve solo lo que de verdad puede usar, sin opciones que confunden.", "En que ayuda:")
nota("en la revision se confirmo ademas que pasar inventario de venta a insumo en Ensambles ya funcionaba para Bodega, Tecnico y Vendedor (no era un error). Todo el uso desde el computador quedo igual, sin cambios.")

# Aclaraciones
part("Aclaraciones", "Cosas que se revisaron y NO eran errores")
lead("estos puntos se reportaron y se revisaron a fondo. Funcionan asi a proposito. Los explicamos para que quede claro el porque.")

item("A", "Cantidades con decimales (por ejemplo 2.5 unidades)")
lead("en todo el sistema las cantidades son numeros enteros (1, 2, 3 y demas), porque el inventario se maneja por unidades completas. El control de stock, los movimientos y las alertas estan armados asi.", "Por que funciona asi:")
lead("el caso de las poleas 2.5 o 3.5 no es una cantidad, es la medida del producto (el tamano de la polea). Eso se resuelve con la correccion del punto 12 (busqueda), no con cantidades decimales. Se venden 1, 2, 3 poleas 2.5, no 2.5 poleas.", "Importante:")
nota("si en el futuro se necesita vender por fraccion (por ejemplo, manguera por metro: 2,5 m), eso requiere una adaptacion especifica del inventario y se puede evaluar como una mejora aparte.")

item("B", "Precios exactos sin redondear (por ejemplo 2.820 pesos)")
lead("en Colombia el punto separa los miles. 2.820 significa dos mil ochocientos veinte pesos, y el sistema lo guarda exacto (2.820), sin redondear. El peso colombiano no usa centavos, por eso los precios van en pesos enteros.", "Por que funciona asi:")
lead("no hay redondeo. Si escribes 2.820, el valor guardado es exactamente $2.820. El punto solo separa los miles, como en cualquier factura.", "En resumen:")

# Glosario
part("Glosario", "Palabras que pueden sonar raras")
tabla(
    ["Palabra", "Que significa"],
    [
        ["Cierre de caja", "Resumen de un periodo (un dia o un rango) con ingresos, egresos y margen."],
        ["Conteo de caja (arqueo)", "Contar el efectivo fisico de la caja y compararlo con lo que el sistema espera."],
        ["Cartera", "Cuentas por cobrar (lo que deben los clientes) y por pagar (lo que se debe a proveedores)."],
        ["Reparacion (orden de trabajo)", "El servicio tecnico o reparacion de un equipo."],
        ["Abono o anticipo", "Un pago parcial que hace el cliente antes de completar el total."],
        ["Caja real", "Contar ingresos y gastos solo cuando el dinero de verdad entra o sale."],
        ["Caja menor", "Gastos pequenos pagados en efectivo (se cuentan como gasto)."],
    ],
)

# Estado
part("Para terminar", "Estado actual y siguiente paso")
bullet("Las correcciones del sistema (por dentro) ya estan aplicadas y funcionando en produccion.")
bullet("Las mejoras de pantalla (incluido todo el uso desde el celular) estan listas y se activan con la proxima publicacion de la aplicacion.")
bullet("Todo el trabajo esta guardado y respaldado en los repositorios del proyecto.")
nota("se reviso a fondo y la caja menor ya se registra correctamente como gasto (tanto en el calculo interno como en la pantalla del cierre, donde aparece en Egresos). No se encontro ningun punto donde la caja menor se sume a las ventas. Si en alguna pantalla se sigue viendo distinto, por favor indicala para revisarla (puede ser informacion guardada en el navegador).", etiqueta="Nota sobre la caja menor:")

build_footer()

out = os.path.join("docs", "Informe-Mejoras-CDV-2026-06-28.docx")
doc.save(out)
print("OK ->", out)
