# -*- coding: utf-8 -*-
"""Genera el Word con las 3 opciones de arquitectura para el rediseño de Órdenes de Trabajo (OT)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

AZUL = RGBColor(0x24, 0x5A, 0x8C)
GRIS = RGBColor(0x55, 0x55, 0x55)
VERDE = RGBColor(0x25, 0x9A, 0x55)
ROJO = RGBColor(0xDC, 0x26, 0x26)

doc = Document()

# Estilos base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h(text, level=1, color=AZUL):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def para(text, bold=False, italic=False, color=None, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def tabla(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ─────────────────────────────────────────────────────────────────────────────
# PORTADA
# ─────────────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Rediseño del módulo de Órdenes de Trabajo (OT)")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = AZUL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Tres opciones de arquitectura para revisión y aprobación")
r.font.size = Pt(13)
r.font.color.rgb = GRIS

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Compresores del Valle S.A.S.")
r.font.size = Pt(11)
r.font.color.rgb = GRIS

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 1. CONTEXTO
# ─────────────────────────────────────────────────────────────────────────────
h("1. ¿Qué es la Orden de Trabajo y qué problema resolvemos?", level=1)
para(
    "La Orden de Trabajo (OT) es el proceso que se usa cuando un cliente trae un equipo "
    "(compresor, cabezote, herramienta neumática, etc.) para que la empresa lo revise, le cotice "
    "un arreglo y lo repare. Hoy la OT está incompleta y genera confusión, por eso este rediseño.",
)
para("Problemas que vamos a corregir:", bold=True)
bullet("Al descargar los repuestos de una OT, esa descarga no se asocia a una venta (la plata de la OT no se ve como ingreso de venta).", )
bullet("Los abonos de la OT no se reflejan donde el equipo los busca, lo que lleva a registrarlos mal (hoy se está creando un 'servicio' llamado abono y facturándolo, lo que infla las ventas con plata que no es venta real).", )
bullet("El perfil de Ventas no puede descargar repuestos en la OT (un problema de permisos).", )
bullet("El flujo no es claro ni ordenado: el usuario tiene que recordar qué hacer y se pierde.", )

para("Objetivo del rediseño:", bold=True, color=VERDE)
bullet("La OT funciona como algo totalmente aparte de Ventas y Cotizaciones: genera su propia cotización y su propia venta.", )
bullet("La descarga de repuestos resta del inventario, y la venta suma a los ingresos.", )
bullet("Cero doble contabilización: cada peso se cuenta una sola vez.", )
bullet("Un flujo guiado con pasos obligatorios: si un paso no se completa, el sistema no deja avanzar.", )

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 2. BASE COMÚN
# ─────────────────────────────────────────────────────────────────────────────
h("2. Base común a las tres opciones", level=1)
para("Estos puntos son iguales en las tres propuestas. Lo único que cambia entre opciones es CUÁNDO y CÓMO la plata se vuelve ingreso (ver sección 3).", italic=True, color=GRIS)

h("2.1 La OT es autocontenida", level=2)
para("Tiene su propio documento, su propia cotización interna y su propia venta. El usuario no entra al módulo de Cotizaciones ni al de Ventas. Pero el dinero sí llega a Ingresos y el inventario sí baja.")

h("2.2 Repuestos a precio de venta", level=2)
para("Hoy los repuestos de la OT se registran al costo. Pasarán a registrarse al precio de venta (lo que realmente paga el cliente). El costo se sigue guardando internamente para poder calcular la utilidad, pero el cliente ve y paga el precio de venta.")

h("2.3 Flujo guiado con 7 pasos obligatorios", level=2)
para("Una sola pantalla con los pasos en orden. No se puede avanzar al siguiente si el actual no está completo. Cada paso indica claramente qué falta.")
tabla(
    ["Paso", "Quién lo hace", "Qué hace", "Requisito para avanzar"],
    [
        ["1. Recepción", "Vendedora", "Registra datos del cliente y del equipo, marca el checklist de cómo llegó e imprime la constancia para el cliente", "Checklist marcado + equipo descrito"],
        ["2. Diagnóstico", "Técnico (asignado)", "Revisa el equipo y escribe qué hay que hacer", "Diagnóstico escrito"],
        ["3. Cotización", "Vendedora", "Arma la cotización interna (repuestos a precio + mano de obra) y la envía al cliente", "Al menos un ítem o mano de obra"],
        ["4. Autorización + anticipo", "Vendedora", "El cliente aprueba y se registra el anticipo del 50%", "Aprobado y anticipo registrado"],
        ["5. Descarga + trabajo", "Técnico / Bodega", "Descarga los repuestos del inventario (el stock baja) y ejecuta la reparación", "Repuestos descargados"],
        ["6. Terminado", "Técnico", "Marca el trabajo como terminado: 'listo para recoger'", "Trabajo realizado escrito"],
        ["7. Recogida → Venta", "Vendedora", "El cliente recoge, paga el saldo y se genera la venta + la factura", "Saldo cubierto"],
    ],
)
para("Nota sobre el checklist: marca cómo llegó el equipo. Lo que falte marcado NO es necesariamente lo que hay que reparar (por ejemplo, el cliente pudo dejar el filtro o la manguera en casa). El checklist es un registro del estado de recepción, no la base de la cotización.", italic=True, color=GRIS)

h("2.4 Principio para no contar la plata dos veces", level=2)
para("Hay dos relojes distintos que hoy se confunden:")
bullet("Caja (efectivo): la plata entra cuando el cliente paga el abono. Debe verse el día que entra, para que el cuadre de caja dé.", bold_prefix="Caja (efectivo): ")
bullet("Venta (ingreso): el documento de venta que suma a Ventas / Ingresos.", bold_prefix="Venta (ingreso): ")
para("La diferencia entre las tres opciones es exactamente cómo se separan estos dos relojes para que un mismo peso nunca se cuente dos veces.")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 3. LAS TRES OPCIONES
# ─────────────────────────────────────────────────────────────────────────────
h("3. Las tres opciones de arquitectura", level=1)

# OPCIÓN A
h("Opción A — Venta única en la recogida (Recomendada)", level=2, color=VERDE)
para("Idea central: los abonos son anticipos (plata recibida que todavía no es venta). En la recogida, la OT genera una sola venta por el valor total, fechada ese día.", bold=True)
para("Flujo completo, desde que llega el cliente:")
bullet("La vendedora recibe el equipo, marca el checklist e imprime la constancia (todavía no hay plata).", )
bullet("Se asigna un técnico, que diagnostica y dice qué se debe hacer.", )
bullet("La vendedora arma y envía la cotización al cliente.", )
bullet("El cliente aprueba; se registra el anticipo del 50%. Ese dinero entra a CAJA como anticipo recibido (no como venta todavía).", )
bullet("El técnico descarga los repuestos del inventario (el stock baja) y hace el trabajo.", )
bullet("El técnico marca 'Terminado'.", )
bullet("El cliente recoge, paga el saldo, y en ese momento se genera UNA venta por el total (repuestos a precio + mano de obra), sin volver a tocar el inventario. Esa venta concilia los anticipos ya pagados + el saldo.", )
para("Cómo se cuenta la plata:")
bullet("Stock: baja en el paso 5 (descarga).", bold_prefix="Stock: ")
bullet("Ingreso/venta: se reconoce una sola vez, en la recogida.", bold_prefix="Ingreso/venta: ")
bullet("Caja: el efectivo se ve cuando entra (los anticipos); en la recogida solo el saldo es efectivo nuevo.", bold_prefix="Caja: ")
para("Ventajas:", bold=True, color=VERDE)
bullet("Cumple exactamente lo solicitado: la OT genera su propia venta y esa venta suma en la recogida.", )
bullet("Un solo documento de venta por OT; reportes limpios; imposible el doble conteo.", )
para("Desventaja:", bold=True, color=ROJO)
bullet("Requiere llevar un pequeño registro de 'anticipos recibidos' (plata que entró pero que aún no es venta).", )

doc.add_paragraph()

# OPCIÓN B
h("Opción B — Ingreso a medida que entra el dinero", level=2, color=AZUL)
para("Idea central: cada abono cuenta como ingreso el día que entra. En la recogida solo se suma el saldo restante.", bold=True)
para("Flujo completo, desde que llega el cliente:")
bullet("Recepción, diagnóstico y cotización: igual que en la Opción A.", )
bullet("El cliente aprueba y abona el 50%. Ese abono cuenta como ingreso ese mismo día.", )
bullet("El técnico descarga los repuestos (el stock baja) y hace el trabajo.", )
bullet("El técnico marca 'Terminado'.", )
bullet("El cliente recoge y paga el saldo; en la recogida se suma solo el saldo restante (lo ya abonado ya había contado).", )
para("Cómo se cuenta la plata:")
bullet("Stock: baja en el paso 5 (descarga).", bold_prefix="Stock: ")
bullet("Ingreso: repartido en el tiempo (abonos cuando entran + saldo en la recogida). Cada peso cuenta una sola vez.", bold_prefix="Ingreso: ")
bullet("Caja: cuadra perfecto, porque el dinero y el ingreso ocurren el mismo día.", bold_prefix="Caja: ")
para("Ventajas:", bold=True, color=VERDE)
bullet("La caja siempre cuadra sin necesidad de registro de anticipos.", )
bullet("Es el cambio más pequeño y simple sobre lo que ya existe.", )
para("Desventaja:", bold=True, color=ROJO)
bullet("El ingreso de una OT queda repartido en varias fechas; no hay un único documento de venta por el total (en la recogida solo 'se suma' el saldo).", )

doc.add_paragraph()

# OPCIÓN C
h("Opción C — OT como cuenta por cobrar (contablemente pura)", level=2, color=AZUL)
para("Idea central: al autorizar (cuando inicia el trabajo) se genera la venta completa de una vez y queda una cuenta por cobrar. Los abonos son pagos contra esa cuenta.", bold=True)
para("Flujo completo, desde que llega el cliente:")
bullet("Recepción, diagnóstico y cotización: igual que en la Opción A.", )
bullet("El cliente aprueba; en ese momento se genera la venta completa por el total y queda una cuenta por cobrar. El stock baja ahí.", )
bullet("El anticipo del 50% y los pagos siguientes abonan esa cuenta por cobrar (no son ingresos nuevos, son pagos).", )
bullet("El técnico hace el trabajo y marca 'Terminado'.", )
bullet("El cliente recoge y paga el saldo de la cuenta por cobrar.", )
para("Cómo se cuenta la plata:")
bullet("Stock: baja al autorizar / descargar.", bold_prefix="Stock: ")
bullet("Ingreso/venta: se reconoce al autorizar (por devengo), no en la recogida.", bold_prefix="Ingreso/venta: ")
bullet("Pagos: los abonos y el saldo se registran como cobranza de la cuenta por cobrar.", bold_prefix="Pagos: ")
para("Ventajas:", bold=True, color=VERDE)
bullet("Es lo más correcto contablemente (la venta se registra cuando se compromete el trabajo).", )
bullet("Reutiliza el módulo de cuentas por cobrar que ya existe en el sistema.", )
para("Desventaja:", bold=True, color=ROJO)
bullet("Contradice la preferencia del negocio: el ingreso entra al autorizar, no en la recogida.", )

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 4. COMPARACIÓN
# ─────────────────────────────────────────────────────────────────────────────
h("4. Comparación rápida", level=1)
tabla(
    ["Criterio", "Opción A", "Opción B", "Opción C"],
    [
        ["¿Cuándo suma el ingreso?", "En la recogida (una sola venta)", "Abonos cuando entran + saldo en recogida", "Al autorizar (devengo)"],
        ["¿Un solo documento de venta?", "Sí", "No (repartido)", "Sí"],
        ["Doble conteo", "Eliminado", "Eliminado", "Eliminado"],
        ["Cuadre de caja", "Necesita registro de anticipos", "Cuadra solo, muy simple", "Necesita módulo de cartera"],
        ["Esfuerzo de implementación", "Medio", "Bajo", "Medio-alto"],
        ["¿Cumple 'venta en la recogida'?", "Sí (total)", "Parcial (solo el saldo)", "No"],
    ],
    widths=[1.7, 1.5, 1.5, 1.5],
)

h("Recomendación", level=2, color=VERDE)
para(
    "Recomendamos la Opción A: es la única que cumple las tres condiciones a la vez — la OT genera "
    "su propia venta, esa venta suma en el momento de la recogida, y nunca hay doble conteo. "
    "La Opción B es una alternativa válida si se prioriza que el cuadre de caja sea lo más simple posible, "
    "a costa de no tener un único documento de venta por OT."
)

# ─────────────────────────────────────────────────────────────────────────────
# 5. CÓMO SE VERÁ (FRONTEND)
# ─────────────────────────────────────────────────────────────────────────────
h("5. Cómo se rediseña la pantalla (igual en las tres opciones)", level=1)
bullet("Se eliminan los botones sueltos de estado. En su lugar, un asistente de 7 pasos en la parte superior de la OT: verde = completo, azul = paso actual, gris = bloqueado. No se puede saltar pasos.", )
bullet("Cada paso es su propio panel, con su acción y un aviso claro de qué falta (por ejemplo: 'Falta registrar el anticipo del 50%').", )
bullet("Cada rol ve resaltado su paso: la vendedora entra y ve 'Recepción' / 'Recogida'; el técnico ve 'Diagnóstico' / 'Trabajo'.", )
bullet("En celular: los pasos se ven como un acordeón vertical, con botones grandes (aptos para usar con guantes).", )
bullet("El paso 'Terminado' incluye el botón 'Convertir a venta / Entregar', con confirmación antes de cerrar.", )
bullet("La impresión sale en dos momentos: constancia de recepción (paso 1, sin abono todavía) y factura final (paso 7, con total, abonos y saldo).", )

doc.add_paragraph()
para("Documento preparado para revisión del cliente. Una vez se elija la opción, se procede al plan de implementación detallado.", italic=True, color=GRIS)

out = "docs/Rediseno-OT-3-Opciones.docx"
doc.save(out)
print("OK ->", out)
